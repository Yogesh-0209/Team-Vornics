import os
import os.path
import fitz  # PyMuPDF
from docx import Document
import tempfile
from typing import Optional, Dict, Any
import logging
import re
import traceback

# Import the AI Extractor
try:
    from ai_extractor import AIExtractor
    ai_available = True
except ImportError:
    ai_available = False
    
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Enhanced document processor with better text extraction and preprocessing
    Integrated with AI capabilities for advanced extraction
    """
    
    def __init__(self):
        self.supported_formats = {
            "application/pdf": self._extract_from_pdf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_from_docx,
            "application/msword": self._extract_from_doc
        }
        
        # Initialize AI extractor if available
        self.ai_extractor = None
        if ai_available:
            try:
                logger.info("Initializing AI Extractor in DocumentProcessor")
                self.ai_extractor = AIExtractor()
                logger.info("AI Extractor initialized successfully")
            except Exception as e:
                logger.error(f"Failed to initialize AI Extractor: {str(e)}")
                logger.error(traceback.format_exc())
    
    def extract_text(self, file_path: str, content_type: str) -> str:
        """
        Extract and preprocess text from document based on content type
        """
        if content_type not in self.supported_formats:
            raise ValueError(f"Unsupported file type: {content_type}")
        
        logger.info(f"Extracting text from {file_path} (type: {content_type})")
        
        extractor = self.supported_formats[content_type]
        raw_text = extractor(file_path)
        
        # Preprocess the extracted text
        processed_text = self.preprocess_text(raw_text)
        
        logger.info(f"Extracted {len(processed_text)} characters of processed text")
        return processed_text
        
    def process_document_with_ai(self, file_path: str, content_type: str, document_name: str) -> Dict[str, Any]:
        """
        Process document using AI extraction capabilities
        
        Args:
            file_path: Path to the document file
            content_type: MIME type of the document
            document_name: Name of the document
            
        Returns:
            Dictionary containing structured data extracted from the document
        """
        logger.info(f"Processing document with AI: {document_name}")
        
        # Check if AI extractor is available
        if not self.ai_extractor or not getattr(self.ai_extractor, 'ai_ready', False):
            logger.warning("AI Extractor not available or not ready")
            return {"error": "AI extraction not available", "ai_extraction": False}
        
        try:
            # Extract text from document
            extracted_text = self.extract_text(file_path, content_type)
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                logger.warning("Extracted text is too short for AI processing")
                return {"error": "Extracted text is too short for AI processing", "ai_extraction": False}
            
            # Process the extracted text with AI
            logger.info(f"Sending {len(extracted_text)} characters to AI Extractor")
            result = self.ai_extractor.extract_structured_data(extracted_text, document_name)
            
            logger.info(f"AI extraction completed with {len(result.get('events', []))} events")
            return result
            
        except Exception as e:
            logger.error(f"Error in AI document processing: {str(e)}")
            logger.error(traceback.format_exc())
            return {"error": f"AI processing failed: {str(e)}", "ai_extraction": False}
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF with enhanced extraction
        """
        try:
            logger.info(f"Opening PDF file: {file_path}")
            # Check if file exists and is readable
            if not os.path.exists(file_path):
                logger.error(f"PDF file does not exist: {file_path}")
                raise FileNotFoundError(f"PDF file not found: {file_path}")
                
            if not os.access(file_path, os.R_OK):
                logger.error(f"PDF file is not readable: {file_path}")
                raise PermissionError(f"Cannot read PDF file: {file_path}")
                
            # Get file size
            file_size = os.path.getsize(file_path)
            logger.info(f"PDF file size: {file_size} bytes")
            
            # Check if file is empty
            if file_size == 0:
                logger.error(f"PDF file is empty: {file_path}")
                raise ValueError(f"PDF file is empty: {file_path}")
            
            # Check if file is a valid PDF (should start with %PDF)
            with open(file_path, 'rb') as f:
                header = f.read(5).decode('ascii', errors='ignore')
                if not header.startswith('%PDF'):
                    logger.error(f"File is not a valid PDF (header: {header}): {file_path}")
                    raise ValueError(f"File is not a valid PDF: {file_path}")
                
            doc = fitz.open(file_path)
            text = ""
            
            logger.info(f"PDF has {len(doc)} pages")
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                logger.info(f"Processing page {page_num+1}")
                
                # Try different extraction methods
                try:
                    page_text = page.get_text()
                except Exception as page_error:
                    logger.error(f"Error extracting text from page {page_num+1}: {str(page_error)}")
                    page_text = ""
                
                # If text extraction fails, try OCR-like extraction
                if not page_text.strip():
                    logger.info(f"Standard text extraction failed for page {page_num+1}, trying block extraction")
                    try:
                        # Get text blocks with positioning info
                        blocks = page.get_text("dict")
                        for block in blocks.get("blocks", []):
                            if "lines" in block:
                                for line in block["lines"]:
                                    for span in line.get("spans", []):
                                        text += span.get("text", "") + " "
                                    text += "\n"
                    except Exception as block_error:
                        logger.error(f"Block extraction failed for page {page_num+1}: {str(block_error)}")
                else:
                    logger.info(f"Extracted {len(page_text)} characters from page {page_num+1}")
                    text += page_text
                
                text += "\n\n"  # Add page separator
            
            doc.close()
            
            if not text.strip():
                logger.error("No text could be extracted from PDF")
                raise Exception("No text could be extracted from PDF")
            
            logger.info(f"Successfully extracted {len(text.strip())} characters from PDF")
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file with enhanced table and formatting handling
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables with better formatting
            for table in doc.tables:
                text += "\n[TABLE]\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "[/TABLE]\n\n"
            
            # Extract headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += "[HEADER] " + paragraph.text + "\n"
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += "[FOOTER] " + paragraph.text + "\n"
            
            if not text.strip():
                raise Exception("No text could be extracted from DOCX")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_doc(self, file_path: str) -> str:
        """
        Extract text from DOC file (legacy Word format)
        """
        try:
            # First try to open as DOCX (sometimes works)
            return self._extract_from_docx(file_path)
        except:
            # For true DOC files, we'd need additional libraries like python-docx2txt
            # For now, provide a helpful error message
            raise Exception(
                "Legacy DOC format detected. Please convert to DOCX or PDF format for processing. "
                "You can use Microsoft Word or LibreOffice to convert the file."
            )
    
    def preprocess_text(self, text: str) -> str:
        """
        Enhanced text preprocessing for better event extraction
        """
        if not text:
            return ""
        
        # Remove excessive whitespace but preserve structure
        lines = []
        for line in text.split('\n'):
            cleaned_line = re.sub(r'\s+', ' ', line.strip())
            if cleaned_line:
                lines.append(cleaned_line)
        
        # Join lines back together
        cleaned_text = '\n'.join(lines)
        
        # Normalize common maritime terms and abbreviations
        cleaned_text = self._normalize_maritime_terms(cleaned_text)
        
        # Fix common OCR errors
        cleaned_text = self._fix_ocr_errors(cleaned_text)
        
        # Normalize date and time formats
        cleaned_text = self._normalize_datetime_formats(cleaned_text)
        
        return cleaned_text
    
    def _normalize_maritime_terms(self, text: str) -> str:
        """
        Normalize common maritime terms and abbreviations
        """
        # Common maritime term normalizations
        normalizations = {
            r'\bvsl\b': 'vessel',
            r'\bmv\b': 'MV',
            r'\bms\b': 'MS',
            r'\bmt\b': 'MT',
            r'\bberth(?:ed|ing)?\b': 'berthed',
            r'\bunberth(?:ed|ing)?\b': 'unberthed',
            r'\banchor(?:ed|ing)?\b': 'anchored',
            r'\bdisch(?:arge|arging)?\b': 'discharge',
            r'\bcomm(?:enced|encing)?\b': 'commenced',
            r'\bcomp(?:leted|leting)?\b': 'completed',
            r'\bhrs?\b': 'hours',
            r'\bmins?\b': 'minutes',
            r'\bsecs?\b': 'seconds'
        }
        
        for pattern, replacement in normalizations.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _fix_ocr_errors(self, text: str) -> str:
        """
        Fix common OCR errors in maritime documents
        """
        # Common OCR error corrections
        corrections = [
            (r'\b0(\d)\b', lambda m: m.group(1)),  # Fix leading zeros in times
            (r'\bl(\d)\b', lambda m: '1' + m.group(1)),  # Fix 'l' mistaken for '1'
            (r'\bO(\d)\b', lambda m: '0' + m.group(1)),  # Fix 'O' mistaken for '0'
            (r'\b(\d)l\b', lambda m: m.group(1) + '1'),  # Fix trailing 'l' in numbers
            (r'\b(\d)O\b', lambda m: m.group(1) + '0')  # Fix trailing 'O' in numbers
        ]
        
        for pattern, replacement in corrections:
            text = re.sub(pattern, replacement, text)
        
        return text
    
    def _normalize_datetime_formats(self, text: str) -> str:
        """
        Normalize various date and time formats found in SoF documents
        """
        # This is a basic implementation - can be expanded based on common formats
        
        # Normalize time formats (e.g., "1430 hrs" -> "14:30")
        text = re.sub(r'\b(\d{1,2})(\d{2})\s*(?:hrs?|hours?)\b', r'\1:\2', text, flags=re.IGNORECASE)
        
        # Normalize date separators
        text = re.sub(r'\b(\d{1,2})[\/\-\.](\d{1,2})[\/\-\.](\d{2,4})\b', r'\1/\2/\3', text)
        
        return text
    
    def extract_metadata(self, file_path: str, content_type: str) -> dict:
        """
        Extract metadata from document
        """
        metadata = {
            "file_size": os.path.getsize(file_path),
            "content_type": content_type,
            "extraction_method": content_type.split('/')[-1]
        }
        
        try:
            if content_type == "application/pdf":
                doc = fitz.open(file_path)
                metadata.update({
                    "page_count": len(doc),
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "creator": doc.metadata.get("creator", ""),
                    "creation_date": doc.metadata.get("creationDate", "")
                })
                doc.close()
            
            elif "wordprocessingml" in content_type:
                doc = Document(file_path)
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                })
        
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")
        
        return metadata
